"""Word提取器，负责从Word文件中提取内容"""
import os
from typing import List, Dict, Any
import docx

from app.core.logger import Logger
from app.rag.extractor.extractor_base import BaseExtractor

class WordExtractor(BaseExtractor):
    """Word提取器，负责从Word文件中提取内容"""
    
    async def extract(self, file_path: str) -> List[Dict[str, Any]]:
        """从Word文件中提取内容
        
        Args:
            file_path: Word文件路径
            
        Returns:
            List[Dict[str, Any]]: 提取的文本内容列表
        """
        import time
        start_time = time.time()
        
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                Logger.error(f"Word文件不存在: {file_path}")
                return []
            
            # 获取文件大小
            file_size = os.path.getsize(file_path)
            Logger.debug(f"开始提取Word文件: {file_path}, 大小: {file_size} bytes")
                
            # 打开Word文件
            open_start_time = time.time()
            doc = docx.Document(file_path)
            open_time = time.time() - open_start_time
            
            Logger.debug(f"Word文件打开成功: 耗时 {open_time:.3f}秒")
            
            # 提取文本内容
            paragraph_start_time = time.time()
            full_text = []
            paragraph_count = 0
            
            for para in doc.paragraphs:
                if para.text.strip():  # 忽略空段落
                    full_text.append(para.text)
                    paragraph_count += 1
                    
            paragraph_time = time.time() - paragraph_start_time
            Logger.debug(f"段落提取完成: {paragraph_count} 个段落, 耗时: {paragraph_time:.3f}秒")
                    
            # 提取表格内容
            table_start_time = time.time()
            table_count = 0
            table_row_count = 0
            
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # 忽略空单元格
                            row_text.append(cell.text.strip())
                    if row_text:  # 忽略空行
                        full_text.append(" | ".join(row_text))
                        table_row_count += 1
                        
            table_time = time.time() - table_start_time
            Logger.debug(f"表格提取完成: {table_count} 个表格, {table_row_count} 行数据, 耗时: {table_time:.3f}秒")
                        
            # 合并文本
            merge_start_time = time.time()
            content = "\n".join(full_text)
            merge_time = time.time() - merge_start_time
            
            # 计算总处理时间
            total_time = time.time() - start_time
            
            # 记录提取成功的详细信息
            Logger.debug(f"Word提取完成: {file_path}")
            Logger.debug(f"  - 段落数: {paragraph_count}")
            Logger.debug(f"  - 表格数: {table_count}")
            Logger.debug(f"  - 表格行数: {table_row_count}")
            Logger.debug(f"  - 文本长度: {len(content)}")
            Logger.debug(f"  - 总耗时: {total_time:.3f}秒")
            
            # 记录性能指标
            Logger.rag_performance_metrics(
                operation="word_extraction",
                duration=total_time,
                file_size=file_size,
                file_type="word",
                content_length=len(content),
                paragraph_count=paragraph_count,
                table_count=table_count,
                table_row_count=table_row_count,
                open_time=open_time,
                paragraph_time=paragraph_time,
                table_time=table_time,
                merge_time=merge_time
            )
            
            # 返回提取的内容
            return [{
                "text": content,
                "metadata": {
                    "source_type": "word",
                    "file_path": file_path,
                    "paragraphs": paragraph_count,
                    "tables": table_count,
                    "table_rows": table_row_count,
                    "content_length": len(content),
                    "extraction_time": total_time
                }
            }]
            
        except Exception as e:
            # 计算处理时间
            total_time = time.time() - start_time
            
            import traceback
            error_info = traceback.format_exc()
            
            Logger.error(f"Word提取失败: {file_path}")
            Logger.error(f"错误信息: {str(e)}")
            Logger.debug(f"堆栈跟踪:\n{error_info}")
            
            # 记录失败的性能指标
            Logger.rag_performance_metrics(
                operation="word_extraction_failed",
                duration=total_time,
                file_size=file_size if 'file_size' in locals() else 0,
                file_type="word",
                error=str(e)
            )
            
            return []